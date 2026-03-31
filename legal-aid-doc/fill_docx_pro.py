#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
法律援助文档填充工具 - Pro版本
支持多个文档模板批量填充 + 智能占位符识别
基于用户的sc.py改进，完整保持原格式

优化1: 增强占位符检测 - 支持多种格式、模糊匹配、多行内容、条件占位符
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_UNDERLINE
import re
import os
import glob
import json
from datetime import datetime
from typing import Dict, List, Optional, Any

# 导入增强的占位符检测模块
try:
    from placeholder_detector import PlaceholderDetector, PlaceholderMatch, MultiLinePlaceholderHandler
    PLACEHOLDER_DETECTOR_AVAILABLE = True
except ImportError:
    PLACEHOLDER_DETECTOR_AVAILABLE = False
    print("警告: placeholder_detector模块未找到，将使用基础占位符检测")

# 导入案件类型管理模块（优化2）
try:
    from case_type_manager import CaseTypeManager, CaseCategory, detect_case_type, get_case_config
    CASE_TYPE_MANAGER_AVAILABLE = True
except ImportError:
    CASE_TYPE_MANAGER_AVAILABLE = False
    print("警告: case_type_manager模块未找到，将使用基础案件类型检测")

# 导入交互式填充模块（优化4）
try:
    from interactive_filler import InteractiveFiller, prompt_for_missing
    INTERACTIVE_FILLER_AVAILABLE = True
except ImportError:
    INTERACTIVE_FILLER_AVAILABLE = False
    print("警告: interactive_filler模块未找到，将使用基础填充模式")

# ==================== 配置 ====================
CHECKBOX_CHARS = ['□', '☐', '☑', '✓', '√', '■', '▢']
SYMBOL_FONTS = ['Wingdings', 'Symbol', 'Arial Unicode MS', 'MS Gothic']

# 默认配置
DEFAULT_INPUT_DIR = r"d:\trae\法律援助文书"
DEFAULT_ELEMENT_FILE = "元素_民事.txt"
DEFAULT_TEMPLATE_PATTERN = "*.docx"
DEFAULT_EXCLUDE_PATTERNS = ["*_已填充.docx", "*_Pro版.docx", "~$*.docx"]

# 智能映射规则 - 根据上下文关键词匹配元素
SMART_MAPPING_RULES = {
    # 表格表头关键词 -> 元素key
    '受援人': 'weitr',
    '案件编号': 'bianh',
    '案由': 'anyou',
    '承办人': None,  # 固定值，不替换
    '办案机关': 'badw',
    '所处阶段': 'jied',
    '指派日期': 'zprq',
    '结案日期': 'jarq',
    '结案时间': 'jarq',
    '时间': 'gcsj{}',  # 需要行号
    '方式': 'gcfs{}',
    '主要内容': 'gcnr{}',
    '身份证号': 'wtrsfz',
    '电话': 'wtrdh',
    '性别': 'wtrxb',
    '出生': 'wtrcs',
    '住址': 'wtrzz',
}

# 元素文件类型映射
CASE_TYPE_MAPPING = {
    '民事': '元素_民事.txt',
    '刑事': '元素_刑事.txt',
    '行政': '元素_行政.txt',
}

# ==================== 智能占位符识别 ====================

def analyze_table_structure(table):
    """分析表格结构，识别可填充的空白单元格"""
    fillable_cells = []
    
    try:
        # 获取表头（第一行）
        if len(table.rows) == 0:
            return fillable_cells
        
        header_row = table.rows[0]
        headers = []
        
        for cell in header_row.cells:
            header_text = cell.text.strip()
            headers.append(header_text)
        
        # 分析数据行
        for row_idx, row in enumerate(table.rows[1:], start=1):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                # 如果单元格为空或只有占位符，记录为可填充
                if not cell_text or cell_text in ['', ' ', '\n']:
                    # 根据列标题确定元素key
                    if col_idx < len(headers):
                        header = headers[col_idx]
                        element_key = None
                        
                        # 匹配规则
                        for keyword, key_template in SMART_MAPPING_RULES.items():
                            if keyword in header:
                                if key_template and '{}' in key_template:
                                    # 需要行号的字段
                                    element_key = key_template.format(row_idx)
                                else:
                                    element_key = key_template
                                break
                        
                        if element_key:
                            fillable_cells.append({
                                'row': row_idx,
                                'col': col_idx,
                                'cell': cell,
                                'header': header,
                                'element_key': element_key
                            })
    except Exception as e:
        print(f"  分析表格时出错: {e}")
    
    return fillable_cells

def detect_context_keywords(paragraph, elements):
    """检测段落中的上下文关键词，智能填充"""
    text = paragraph.text.strip()
    
    # 如果段落已有内容，不处理
    if len(text) > 5:
        return None
    
    # 检测关键词
    for keyword, element_key in SMART_MAPPING_RULES.items():
        if keyword in text and element_key and '{}' not in element_key:
            if element_key in elements:
                return {
                    'keyword': keyword,
                    'element_key': element_key,
                    'value': elements[element_key]
                }
    
    return None

# ==================== 格式处理函数 ====================

def get_complete_run_format(run):
    """完整获取run的所有格式属性"""
    if not run or not hasattr(run, 'font'):
        return {}
    
    try:
        format_info = {
            'font_size': run.font.size,
            'font_name': None,
            'bold': run.bold,
            'italic': run.italic,
            'underline': getattr(run.font, 'underline', None),
            'underline_color': getattr(run.font, 'underline_color', None),
            'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
            'highlight_color': getattr(run.font, 'highlight_color', None),
            'all_caps': getattr(run.font, 'all_caps', None),
            'small_caps': getattr(run.font, 'small_caps', None),
            'strike': getattr(run.font, 'strike', None),
            'double_strike': getattr(run.font, 'double_strike', None),
            'subscript': getattr(run.font, 'subscript', None),
            'superscript': getattr(run.font, 'superscript', None),
            'character_spacing': getattr(run.font, 'character_spacing', None),
            'scaling': getattr(run.font, 'scaling', None),
        }
        
        # 获取字体名称
        try:
            if run.font.name:
                format_info['font_name'] = run.font.name
            elif hasattr(run, '_element'):
                rpr = run._element.rPr
                if rpr is not None:
                    rFonts = rpr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        for attr in ['w:ascii', 'w:eastAsia', 'w:hAnsi', 'w:cs']:
                            font_name = rFonts.get(qn(attr))
                            if font_name:
                                format_info['font_name'] = font_name
                                break
        except Exception:
            pass
            
        return format_info
    except Exception:
        return {}

def apply_complete_format(run, format_info):
    """完整应用格式"""
    if not run or not format_info:
        return
    
    try:
        # 设置字体
        if format_info.get('font_name'):
            try:
                run.font.name = format_info['font_name']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), format_info['font_name'])
            except Exception:
                try:
                    rpr = run._element.get_or_add_rPr()
                    rFonts = rpr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rpr.append(rFonts)
                    rFonts.set(qn('w:ascii'), format_info['font_name'])
                    rFonts.set(qn('w:eastAsia'), format_info['font_name'])
                    rFonts.set(qn('w:hAnsi'), format_info['font_name'])
                    rFonts.set(qn('w:cs'), format_info['font_name'])
                except Exception:
                    pass
        
        if format_info.get('font_size'):
            try:
                run.font.size = format_info['font_size']
            except:
                pass
        
        if 'bold' in format_info:
            run.bold = format_info['bold']
        if 'italic' in format_info:
            run.italic = format_info['italic']
        
        if format_info.get('underline') is not None:
            try:
                run.font.underline = format_info['underline']
                if format_info.get('underline_color'):
                    run.font.underline_color = format_info['underline_color']
            except:
                pass
        
        if format_info.get('color'):
            try:
                run.font.color.rgb = format_info['color']
            except:
                pass
        
        format_attrs = ['all_caps', 'small_caps', 'strike', 'double_strike',
                       'subscript', 'superscript', 'character_spacing', 'scaling']
        
        for attr in format_attrs:
            if attr in format_info and format_info[attr] is not None:
                try:
                    setattr(run.font, attr, format_info[attr])
                except:
                    pass
    except Exception:
        pass

def is_checkbox_char(char):
    """检查字符是否为多选框"""
    return char in CHECKBOX_CHARS

def set_table_row_allow_break(row):
    """设置表格行允许跨页断行"""
    try:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        
        cantSplit_elements = trPr.findall(qn('w:cantSplit'))
        for element in cantSplit_elements:
            trPr.remove(element)
        
        existing_heights = trPr.findall(qn('w:trHeight'))
        if not existing_heights:
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), "500")
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)
        else:
            for height_elem in existing_heights:
                height_elem.set(qn('w:hRule'), "atLeast")
    except Exception:
        pass

# ==================== 段落处理 ====================

def process_paragraph(paragraph, sorted_keys, mapping, smart_fill=False, elements=None, use_enhanced=False):
    """
    处理单个段落，执行替换并保持格式
    
    Args:
        use_enhanced: 是否使用增强占位符检测（优化1）
    """
    if not paragraph.text or not paragraph.runs:
        return
    
    # 使用增强检测（优化1）
    if use_enhanced and PLACEHOLDER_DETECTOR_AVAILABLE:
        enhanced_processor.process_paragraph_with_enhanced_detection(paragraph, mapping)
        return
    
    # 智能填充模式
    if smart_fill and elements:
        smart_result = detect_context_keywords(paragraph, elements)
        if smart_result:
            # 在段落末尾添加内容
            run = paragraph.add_run(smart_result['value'])
            return
    
    # 记录每个字符的格式
    char_formats = []
    full_text_chars = []
    
    for run in paragraph.runs:
        if run.text:
            run_format = get_complete_run_format(run)
            for char in run.text:
                format_info = run_format.copy()
                format_info['is_checkbox'] = is_checkbox_char(char)
                char_formats.append(format_info)
                full_text_chars.append(char)
    
    full_text = ''.join(full_text_chars)
    if not full_text.strip():
        return
    
    # 查找所有替换位置
    replacements = []
    for key in sorted_keys:
        pattern = re.compile(re.escape(key))
        for match in pattern.finditer(full_text):
            start, end = match.span()
            value = str(mapping.get(key, ''))
            replacements.append({
                'start': start,
                'end': end,
                'replacement': value
            })
    
    if not replacements:
        return
    
    # 从后往前替换
    replacements.sort(key=lambda x: x['start'], reverse=True)
    
    new_text_chars = list(full_text_chars)
    new_char_formats = list(char_formats)
    
    for rep in replacements:
        start, end = rep['start'], rep['end']
        replacement = rep['replacement']
        
        # 检查是否包含多选框
        contains_checkbox = any(
            new_char_formats[i].get('is_checkbox', False)
            for i in range(start, min(end, len(new_char_formats)))
        )
        
        if contains_checkbox:
            # 特殊处理：保留多选框
            original_chars = new_text_chars[start:end]
            replacement_chars = list(replacement)
            
            new_segment = []
            new_format_segment = []
            
            orig_idx = 0
            repl_idx = 0
            
            while orig_idx < len(original_chars) and repl_idx < len(replacement_chars):
                if is_checkbox_char(original_chars[orig_idx]):
                    new_segment.append(original_chars[orig_idx])
                    new_format_segment.append(new_char_formats[start + orig_idx])
                    orig_idx += 1
                else:
                    if repl_idx < len(replacement_chars):
                        new_segment.append(replacement_chars[repl_idx])
                        format_idx = min(start + orig_idx, len(new_char_formats) - 1)
                        new_format_segment.append(new_char_formats[format_idx])
                        repl_idx += 1
                        orig_idx += 1
            
            while repl_idx < len(replacement_chars):
                new_segment.append(replacement_chars[repl_idx])
                new_format_segment.append(new_char_formats[min(end-1, len(new_char_formats)-1)])
                repl_idx += 1
            
            new_text_chars[start:end] = new_segment
            new_char_formats[start:end] = new_format_segment
        else:
            # 普通替换
            replacement_format = new_char_formats[start] if start < len(new_char_formats) else {}
            new_text_chars[start:end] = list(replacement)
            
            if len(replacement) <= (end - start):
                new_formats_segment = new_char_formats[start:start + len(replacement)]
            else:
                new_formats_segment = new_char_formats[start:end]
                if new_formats_segment:
                    last_format = new_formats_segment[-1]
                    new_formats_segment.extend([last_format] * (len(replacement) - len(new_formats_segment)))
                else:
                    new_formats_segment = [replacement_format] * len(replacement)
            
            new_char_formats[start:end] = new_formats_segment
    
    # 重建段落
    new_full_text = ''.join(new_text_chars)
    paragraph.clear()
    
    if not new_full_text:
        return
    
    # 按格式变化分段创建run
    current_format = new_char_formats[0] if new_char_formats else {}
    current_segment = [new_full_text[0]]
    
    for i in range(1, len(new_full_text)):
        current_char = new_full_text[i]
        current_char_format = new_char_formats[i] if i < len(new_char_formats) else current_format
        
        format_changed = (
            current_format.get('bold') != current_char_format.get('bold') or
            current_format.get('italic') != current_char_format.get('italic') or
            current_format.get('underline') != current_char_format.get('underline') or
            current_format.get('color') != current_char_format.get('color') or
            current_format.get('font_size') != current_char_format.get('font_size') or
            current_format.get('font_name') != current_char_format.get('font_name')
        )
        
        if format_changed:
            if current_segment:
                run = paragraph.add_run(''.join(current_segment))
                apply_complete_format(run, current_format)
                if any(is_checkbox_char(char) for char in current_segment):
                    try:
                        for font in SYMBOL_FONTS:
                            try:
                                run.font.name = font
                                break
                            except:
                                continue
                    except:
                        pass
                current_segment = []
        
        current_segment.append(current_char)
        current_format = current_char_format
    
    if current_segment:
        run = paragraph.add_run(''.join(current_segment))
        apply_complete_format(run, current_format)
        if any(is_checkbox_char(char) for char in current_segment):
            try:
                for font in SYMBOL_FONTS:
                    try:
                        run.font.name = font
                        break
                    except:
                        continue
            except:
                pass

def process_table(table, sorted_keys, mapping, smart_fill=False, elements=None, use_enhanced=False):
    """
    处理表格
    
    Args:
        use_enhanced: 是否使用增强占位符检测（优化1）
    """
    try:
        # 智能填充空白单元格
        if smart_fill and elements:
            fillable_cells = analyze_table_structure(table)
            for cell_info in fillable_cells:
                cell = cell_info['cell']
                element_key = cell_info['element_key']
                
                if element_key in elements:
                    # 填充单元格
                    if cell.paragraphs:
                        p = cell.paragraphs[0]
                        if not p.text.strip() or p.text.strip() == '':
                            p.clear()
                            run = p.add_run(elements[element_key])
        
        # 常规替换处理
        for row in table.rows:
            set_table_row_allow_break(row)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph, sorted_keys, mapping, smart_fill, elements, use_enhanced)
                for nested_table in cell.tables:
                    process_table(nested_table, sorted_keys, mapping, smart_fill, elements, use_enhanced)
    except Exception:
        pass

def process_document(document, mapping, smart_fill=False, elements=None, use_enhanced=False):
    """
    处理整个文档
    
    Args:
        use_enhanced: 是否使用增强占位符检测（优化1）
    """
    if not mapping:
        return document
    
    sorted_keys = sorted(mapping.keys(), key=len, reverse=True)
    
    # 处理段落
    for paragraph in document.paragraphs:
        try:
            process_paragraph(paragraph, sorted_keys, mapping, smart_fill, elements, use_enhanced)
        except Exception:
            pass
    
    # 处理表格
    for table in document.tables:
        try:
            process_table(table, sorted_keys, mapping, smart_fill, elements, use_enhanced)
        except Exception:
            pass
    
    # 处理页眉页脚
    for section in document.sections:
        try:
            for paragraph in section.header.paragraphs:
                process_paragraph(paragraph, sorted_keys, mapping, smart_fill, elements, use_enhanced)
            for paragraph in section.footer.paragraphs:
                process_paragraph(paragraph, sorted_keys, mapping, smart_fill, elements, use_enhanced)
        except Exception:
            pass
    
    return document

# ==================== 元素解析 ====================

def parse_element_file(file_path):
    """解析元素文件"""
    elements = {}
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    sections = re.split(r'\n--', content)
    for section in sections:
        section = section.strip()
        if not section:
            continue
        lines = section.split('\n')
        for line in lines:
            line = line.strip()
            if '::' in line and not line.startswith('#'):
                key, value = line.split('::', 1)
                key = key.strip()
                value = value.strip()
                # 处理转义字符
                value = value.replace('\\n', '\n').replace('\\t', '\t')
                elements[key] = value
    return elements

def build_mapping(elements, category=None):
    """
    构建替换映射
    
    Args:
        category: 案件类型（优化2）
    """
    # 如果启用了案件类型管理器，使用类型特定的映射
    if CASE_TYPE_MANAGER_AVAILABLE and category:
        manager = CaseTypeManager()
        return manager.build_mapping_for_category(elements, category)
    
    # 基础映射（向后兼容）
    mapping = {}
    
    # 基本字段
    basic_keys = ['bianh', 'leib', 'anyou', 'weitr', 'dfdsr', 'badw', 'jied', 
                  'zprq', 'jarq', 'cbxj', 'ljsm', 'gdrq', 'yjrq', 'wtrsfz', 
                  'wtrdh', 'wtrxb', 'wtrcs', 'wtrzz', 'dcdw', 'bdqxx', 'thsj1',
                  'thdd1', 'ajqk', 'ssqq', 'basl', 'ktsj', 'ktdd', 'cbfg', 'sjy']
    
    for key in basic_keys:
        if key in elements:
            mapping[key] = elements[key]
    
    # 办案过程字段
    for i in range(1, 10):
        for prefix in ['gcsj', 'gcfs', 'gcnr', 'thsj', 'thdd']:
            key = f"{prefix}{i}"
            if key in elements:
                mapping[key] = elements[key]
    
    return mapping

# ==================== 增强占位符处理（优化1）====================

class EnhancedPlaceholderProcessor:
    """增强型占位符处理器 - 使用新的检测模块"""
    
    def __init__(self):
        self.detector = PlaceholderDetector() if PLACEHOLDER_DETECTOR_AVAILABLE else None
        self.multiline_handler = MultiLinePlaceholderHandler() if PLACEHOLDER_DETECTOR_AVAILABLE else None
    
    def detect_and_replace(self, text: str, values: Dict[str, Any], 
                          keep_unmatched: bool = True) -> str:
        """
        检测并替换文本中的占位符
        
        Args:
            text: 原始文本
            values: 键值映射
            keep_unmatched: 是否保留未匹配的占位符
        
        Returns:
            替换后的文本
        """
        if not self.detector:
            # 回退到基础替换
            return self._basic_replace(text, values)
        
        return self.detector.replace_placeholders(text, values, keep_unmatched)
    
    def _basic_replace(self, text: str, values: Dict[str, Any]) -> str:
        """基础替换（当增强模块不可用时使用）"""
        result = text
        for key, value in values.items():
            # 支持多种占位符格式
            patterns = [
                f'{{{key}}}',      # {key}
                f'{{{{{key}}}}}',  # {{key}}
                f'[{key}]',        # [key]
                f'[[{key}]]',      # [[key]]
                f'【{key}】',      # 【key】
            ]
            for pattern in patterns:
                result = result.replace(pattern, str(value))
        return result
    
    def extract_required_keys(self, text: str) -> set:
        """提取文本中需要的所有键名"""
        if not self.detector:
            # 基础提取
            keys = set()
            for match in re.finditer(r'\{([^}]+)\}', text):
                keys.add(match.group(1))
            return keys
        
        return self.detector.extract_keys(text)
    
    def validate_template(self, text: str, available_keys: set) -> tuple:
        """
        验证模板中的占位符
        
        Returns:
            (有效键集合, 缺失键集合)
        """
        if not self.detector:
            found = self.extract_required_keys(text)
            valid = found & available_keys
            missing = found - available_keys
            return valid, missing
        
        return self.detector.validate_keys(text, available_keys)
    
    def process_paragraph_with_enhanced_detection(self, paragraph, values: Dict[str, Any]):
        """
        使用增强检测处理段落
        保持原有格式的同时支持多种占位符格式，也支持纯文本键名
        """
        if not paragraph.text or not paragraph.runs:
            return
        
        full_text = paragraph.text
        matches = []
        
        # 使用增强检测器查找格式化的占位符
        if self.detector:
            matches = self.detector.detect_all(full_text)
        
        # 同时支持纯文本键名（如模板中的 weitr, bianh）
        for key in values.keys():
            # 查找纯文本键名（不在其他词中间）
            pattern = r'(?<![a-zA-Z0-9_])' + re.escape(key) + r'(?![a-zA-Z0-9_])'
            for m in re.finditer(pattern, full_text):
                # 检查这个位置是否已经被格式化占位符匹配
                pos = m.start()
                length = m.end() - m.start()
                overlap = False
                for existing in matches:
                    if pos < existing.end_pos and pos + length > existing.start_pos:
                        overlap = True
                        break
                if not overlap:
                    from placeholder_detector import PlaceholderMatch, PlaceholderType
                    matches.append(PlaceholderMatch(
                        original=key,
                        key=key,
                        start_pos=pos,
                        end_pos=m.end(),
                        placeholder_type=PlaceholderType.PLAIN,
                        format_spec=None,
                        default_value=None
                    ))
        
        if not matches:
            return
        
        # 按位置排序（从后往前，避免位置偏移）
        matches.sort(key=lambda x: x.start_pos, reverse=True)
        
        # 直接在原始文本上替换，保留原有格式
        full_text = paragraph.text
        
        # 从后往前替换，避免位置偏移
        for match in matches:
            key = match.key
            value = values.get(key)
            
            # 处理默认值
            if value is None and match.default_value:
                value = match.default_value
            
            if value is None:
                continue
            
            # 应用格式说明
            if match.format_spec and self.detector:
                value = self.detector._apply_format(value, match.format_spec)
            
            replacement = str(value)
            start, end = match.start_pos, match.end_pos
            
            # 找到包含这个占位符的 run
            current_pos = 0
            for run in paragraph.runs:
                run_len = len(run.text)
                run_start = current_pos
                run_end = current_pos + run_len
                
                # 检查这个 run 是否包含占位符
                if run_start <= start < run_end:
                    # 占位符在这个 run 中
                    run_text = run.text
                    local_start = start - run_start
                    local_end = end - run_start
                    
                    # 替换文本
                    new_text = run_text[:local_start] + replacement + run_text[local_end:]
                    run.text = new_text
                    break
                
                current_pos += run_len
    
    def _format_changed(self, fmt1: dict, fmt2: dict) -> bool:
        """检查格式是否变化"""
        keys = ['bold', 'italic', 'underline', 'color', 'font_size', 'font_name']
        return any(fmt1.get(k) != fmt2.get(k) for k in keys)


# 全局处理器实例
enhanced_processor = EnhancedPlaceholderProcessor()

# ==================== 文件处理 ====================

def should_exclude_file(filename, exclude_patterns):
    """检查文件是否应该被排除"""
    for pattern in exclude_patterns:
        if re.match(pattern.replace('*', '.*').replace('?', '.'), filename):
            return True
    return False

def get_template_files(input_dir, pattern=DEFAULT_TEMPLATE_PATTERN, exclude_patterns=None):
    """获取所有模板文件"""
    if exclude_patterns is None:
        exclude_patterns = DEFAULT_EXCLUDE_PATTERNS
    
    template_files = []
    search_pattern = os.path.join(input_dir, pattern)
    
    for file_path in glob.glob(search_pattern):
        filename = os.path.basename(file_path)
        if not should_exclude_file(filename, exclude_patterns):
            template_files.append(file_path)
    
    return sorted(template_files)

def generate_output_path(template_path, suffix="_已填充"):
    """生成输出文件路径"""
    dir_name = os.path.dirname(template_path)
    base_name = os.path.splitext(os.path.basename(template_path))[0]
    output_path = os.path.join(dir_name, f"{base_name}{suffix}.docx")
    return output_path

# ==================== 主函数 ====================

def fill_single_document(template_path, element_file, output_path=None, suffix="_已填充", 
                         smart_fill=False, use_enhanced=False, category=None):
    """
    填充单个文档
    
    Args:
        use_enhanced: 是否使用增强占位符检测（优化1）
        category: 案件类型（优化2）
    """
    
    if output_path is None:
        output_path = generate_output_path(template_path, suffix)
    
    print(f"\n处理: {os.path.basename(template_path)}")
    
    # 解析元素
    elements = parse_element_file(element_file)
    
    # 自动检测案件类型（优化2）
    if category is None and CASE_TYPE_MANAGER_AVAILABLE:
        input_dir = os.path.dirname(element_file)
        detected_category, detected_file = detect_case_type(input_dir, element_file)
        category = detected_category
        print(f"  案件类型: {category.value}")
    
    mapping = build_mapping(elements, category)
    
    if not mapping:
        print(f"  警告: 没有构建到替换规则")
        return False
    
    print(f"  替换规则: {len(mapping)} 个")
    if smart_fill:
        print(f"  智能填充: 已启用")
    if use_enhanced:
        print(f"  增强检测: 已启用（支持多种占位符格式）")
    
    # 处理文档
    document = Document(template_path)
    document = process_document(document, mapping, smart_fill, elements, use_enhanced)
    
    # 保存
    document.save(output_path)
    print(f"  生成: {os.path.basename(output_path)}")
    
    return True

def fill_multiple_documents(input_dir=None, element_file=None, template_pattern=None, 
                            suffix="_已填充", smart_fill=False, use_enhanced=False, 
                            auto_detect_type=True, interactive=False):
    """
    批量填充多个文档
    
    Args:
        use_enhanced: 是否使用增强占位符检测（优化1）
        auto_detect_type: 是否自动检测案件类型（优化2）
        interactive: 是否启用交互式填充（优化4）
    """
    
    if input_dir is None:
        input_dir = DEFAULT_INPUT_DIR
    if element_file is None:
        element_file = os.path.join(input_dir, DEFAULT_ELEMENT_FILE)
    if template_pattern is None:
        template_pattern = DEFAULT_TEMPLATE_PATTERN
    
    # 自动检测案件类型（优化2）
    category = None
    if auto_detect_type and CASE_TYPE_MANAGER_AVAILABLE:
        category, detected_file = detect_case_type(input_dir, element_file)
        if detected_file != element_file and os.path.exists(detected_file):
            element_file = detected_file
    
    print(f"=" * 60)
    print(f"法律援助文档批量填充工具")
    print(f"=" * 60)
    if category:
        config = get_case_config(category)
        print(f"案件类型: {config.display_name}")
    print(f"元素文件: {element_file}")
    print(f"模板目录: {input_dir}")
    print(f"模板匹配: {template_pattern}")
    print(f"智能填充: {'已启用' if smart_fill else '已禁用'}")
    print(f"增强检测: {'已启用' if use_enhanced else '已禁用'} (支持 {{key}}, [[key]], 【key】等格式)")
    print(f"交互模式: {'已启用' if interactive else '已禁用'} (优化4)")
    print(f"=" * 60)
    
    # 检查元素文件
    if not os.path.exists(element_file):
        print(f"错误: 元素文件不存在: {element_file}")
        return
    
    # 获取所有模板文件
    template_files = get_template_files(input_dir, template_pattern)
    
    if not template_files:
        print(f"警告: 未找到匹配的模板文件")
        return
    
    print(f"找到 {len(template_files)} 个模板文件:")
    for i, tf in enumerate(template_files, 1):
        print(f"  {i}. {os.path.basename(tf)}")
    
    print(f"\n开始处理...")
    
    # 解析元素（只解析一次）
    elements = parse_element_file(element_file)
    print(f"解析到 {len(elements)} 个元素")
    
    # 交互式填充缺失字段（优化4）
    missing_fields = []
    if category and CASE_TYPE_MANAGER_AVAILABLE:
        manager = CaseTypeManager()
        is_valid, missing_fields = manager.validate_elements(elements, category)
        if not is_valid:
            print(f"\n⚠️  警告: 以下必填字段缺失: {', '.join(missing_fields)}")
            if interactive and INTERACTIVE_FILLER_AVAILABLE:
                elements = prompt_for_missing(elements, missing_fields)
                # 重新验证
                is_valid, missing_fields = manager.validate_elements(elements, category)
    
    mapping = build_mapping(elements, category)
    print(f"构建 {len(mapping)} 个替换规则")
    
    # 处理每个模板
    success_count = 0
    for template_path in template_files:
        try:
            output_path = generate_output_path(template_path, suffix)
            
            print(f"\n处理: {os.path.basename(template_path)}")
            document = Document(template_path)
            document = process_document(document, mapping, smart_fill, elements, use_enhanced)
            document.save(output_path)
            
            print(f"  ✓ 生成: {os.path.basename(output_path)}")
            success_count += 1
        except Exception as e:
            print(f"  ✗ 失败: {e}")
    
    print(f"\n{'=' * 60}")
    print(f"处理完成: {success_count}/{len(template_files)} 个文件")
    print(f"{'=' * 60}")
    
    return success_count

def main():
    """主函数 - 支持命令行参数"""
    import sys
    
    # 默认配置
    input_dir = DEFAULT_INPUT_DIR
    element_file = os.path.join(input_dir, DEFAULT_ELEMENT_FILE)
    smart_fill = False
    use_enhanced = False  # 优化1: 增强占位符检测
    interactive = False  # 优化4: 交互式填充
    
    # 解析命令行参数
    if len(sys.argv) > 1:
        # 如果提供了参数，使用指定目录
        input_dir = sys.argv[1]
        element_file = os.path.join(input_dir, DEFAULT_ELEMENT_FILE)
    
    if len(sys.argv) > 2:
        # 指定元素文件（支持相对路径或绝对路径）
        arg2 = sys.argv[2]
        if os.path.isabs(arg2) or os.path.exists(arg2):
            # 绝对路径或已存在的文件
            element_file = arg2
        else:
            # 相对路径，拼接输入目录
            element_file = os.path.join(input_dir, arg2)
    
    if len(sys.argv) > 3:
        # 启用智能填充、增强检测或交互模式
        arg = sys.argv[3].lower()
        smart_fill = arg in ['true', '1', 'yes', 'smart']
        use_enhanced = arg in ['enhanced', 'pro', 'advanced', 'all']
        interactive = arg in ['interactive', 'i', '对话', '交互']
    
    if len(sys.argv) > 4:
        # 单独控制增强检测或交互模式
        arg = sys.argv[4].lower()
        use_enhanced = arg in ['true', '1', 'yes', 'enhanced']
        interactive = arg in ['interactive', 'i', '对话', '交互']
    
    # 执行批量填充
    fill_multiple_documents(input_dir, element_file, smart_fill=smart_fill, 
                           use_enhanced=use_enhanced, interactive=interactive)

if __name__ == "__main__":
    main()

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word模板处理器
支持读取、解析和填充Word格式的法律文书模板
"""

import argparse
import json
import sys
import re
from pathlib import Path
from typing import Dict, List, Tuple, Any
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.table import Table
except ImportError:
    print("错误：未安装 python-docx 库，请运行: pip install python-docx")
    sys.exit(1)


class WordTemplateProcessor:
    """Word模板处理器"""
    
    def __init__(self):
        """初始化处理器"""
        self.placeholder_pattern = re.compile(r'\{\{(\w+)\}\}|\[(.+?)\]|【(.+?)】')
    
    def read_template(self, template_path: str) -> Dict:
        """
        读取Word模板文件
        
        Args:
            template_path: Word模板文件路径
            
        Returns:
            模板结构信息
        """
        template_path = Path(template_path)
        
        if not template_path.exists():
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        if not template_path.suffix.lower() in ['.docx', '.doc']:
            raise ValueError(f"不支持的文件格式: {template_path.suffix}，仅支持.docx和.doc格式")
        
        # 读取Word文档
        doc = Document(str(template_path))
        
        # 提取模板信息
        template_info = {
            'file_name': template_path.name,
            'file_path': str(template_path),
            'title': self._extract_title(doc),
            'paragraphs': [],
            'tables': [],
            'placeholders': set(),
            'structure': {}
        }
        
        # 提取段落内容和占位符
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                # 提取占位符
                placeholders = self._extract_placeholders(text)
                template_info['placeholders'].update(placeholders)
                
                template_info['paragraphs'].append({
                    'index': i,
                    'text': text,
                    'placeholders': list(placeholders),
                    'style': para.style.name if para.style else 'Normal'
                })
        
        # 提取表格
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
                
                # 提取表格中的占位符
                for cell_text in row_data:
                    placeholders = self._extract_placeholders(cell_text)
                    template_info['placeholders'].update(placeholders)
            
            template_info['tables'].append({
                'index': i,
                'data': table_data,
                'rows': len(table.rows),
                'cols': len(table.columns)
            })
        
        # 转换占位符集合为列表
        template_info['placeholders'] = list(template_info['placeholders'])
        
        # 分析模板结构
        template_info['structure'] = self._analyze_structure(template_info)
        
        return template_info
    
    def _extract_title(self, doc: Document) -> str:
        """提取文档标题"""
        if doc.paragraphs:
            # 通常第一个段落是标题
            first_para = doc.paragraphs[0].text.strip()
            if first_para:
                return first_para
        return "未命名模板"
    
    def _extract_placeholders(self, text: str) -> set:
        """
        提取文本中的占位符
        
        Args:
            text: 文本内容
            
        Returns:
            占位符集合
        """
        placeholders = set()
        
        # 匹配 {{placeholder}} 格式
        matches = self.placeholder_pattern.findall(text)
        for match in matches:
            # match是一个元组，取非空的值
            for m in match:
                if m:
                    placeholders.add(m.strip())
        
        # 匹配下划线占位符（如：_______）
        underline_pattern = re.compile(r'_{3,}')
        if underline_pattern.search(text):
            placeholders.add('需要填充的内容')
        
        return placeholders
    
    def _analyze_structure(self, template_info: Dict) -> Dict:
        """
        分析模板结构
        
        Args:
            template_info: 模板信息
            
        Returns:
            结构分析结果
        """
        structure = {
            'type': 'unknown',
            'sections': [],
            'has_table': len(template_info['tables']) > 0
        }
        
        # 判断文书类型
        title = template_info['title']
        if '起诉状' in title:
            structure['type'] = '诉讼文书'
        elif '合同' in title:
            structure['type'] = '合同协议'
        elif '答辩状' in title:
            structure['type'] = '诉讼文书'
        elif '申请书' in title:
            structure['type'] = '申请文书'
        
        # 提取章节
        for para in template_info['paragraphs']:
            text = para['text']
            # 识别标题行（通常包含【】或第X条等）
            if text.startswith('【') or text.startswith('第') or text.endswith('：'):
                structure['sections'].append(text)
        
        return structure
    
    def fill_template(self, template_path: str, data: Dict, output_path: str) -> str:
        """
        填充Word模板
        
        Args:
            template_path: 模板文件路径
            data: 填充数据字典
            output_path: 输出文件路径
            
        Returns:
            生成的文件路径
        """
        template_path = Path(template_path)
        output_path = Path(output_path)
        
        # 确保输出目录存在
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 读取模板
        doc = Document(str(template_path))
        
        # 填充段落
        for para in doc.paragraphs:
            self._fill_paragraph(para, data)
        
        # 填充表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._fill_paragraph(para, data)
        
        # 保存文档
        doc.save(str(output_path))
        
        print(f"文书已生成: {output_path}")
        return str(output_path)
    
    def _fill_paragraph(self, para, data: Dict):
        """
        填充段落中的占位符
        
        Args:
            para: 段落对象
            data: 填充数据
        """
        text = para.text
        
        # 替换 {{placeholder}} 格式
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in text:
                text = text.replace(placeholder, str(value))
        
        # 替换 [placeholder] 格式
        for key, value in data.items():
            placeholder = f"[{key}]"
            if placeholder in text:
                text = text.replace(placeholder, str(value))
        
        # 替换 【placeholder】 格式
        for key, value in data.items():
            placeholder = f"【{key}】"
            if placeholder in text:
                text = text.replace(placeholder, str(value))
        
        # 替换下划线（如果数据中有对应字段）
        if '_______' in text or '______' in text or '_____' in text:
            # 获取段落中第一个非空数据值
            for key, value in data.items():
                if value and str(value).strip():
                    text = re.sub(r'_{3,}', str(value), text, count=1)
                    break
        
        # 如果文本有变化，更新段落
        if text != para.text:
            # 清空段落
            para.clear()
            # 添加新文本
            run = para.add_run(text)
    
    def extract_template_info(self, template_path: str, output_json: str = None) -> Dict:
        """
        提取模板信息并生成JSON描述
        
        Args:
            template_path: 模板文件路径
            output_json: 输出JSON文件路径（可选）
            
        Returns:
            模板信息字典
        """
        template_info = self.read_template(template_path)
        
        # 生成要素清单
        elements = []
        for placeholder in template_info['placeholders']:
            elements.append({
                'name': placeholder,
                'required': True,
                'description': f"需要填写的{placeholder}信息"
            })
        
        result = {
            'id': Path(template_path).stem,
            'name': template_info['title'],
            'source_file': template_path,
            'type': template_info['structure']['type'],
            'placeholders': template_info['placeholders'],
            'elements': elements,
            'paragraph_count': len(template_info['paragraphs']),
            'table_count': len(template_info['tables']),
            'sections': template_info['structure']['sections'],
            'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        
        # 保存为JSON
        if output_json:
            with open(output_json, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
            print(f"模板信息已导出: {output_json}")
        
        return result
    
    def batch_fill(self, template_path: str, data_list: List[Dict], output_dir: str) -> List[str]:
        """
        批量填充模板
        
        Args:
            template_path: 模板文件路径
            data_list: 数据列表
            output_dir: 输出目录
            
        Returns:
            生成的文件路径列表
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        generated_files = []
        template_name = Path(template_path).stem
        
        for i, data in enumerate(data_list, 1):
            # 生成输出文件名
            output_file = output_dir / f"{template_name}_{i}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
            
            # 填充模板
            file_path = self.fill_template(template_path, data, str(output_file))
            generated_files.append(file_path)
        
        print(f"批量生成完成，共生成 {len(generated_files)} 份文书")
        return generated_files
    
    def compare_templates(self, template_path1: str, template_path2: str) -> Dict:
        """
        比较两个模板的差异
        
        Args:
            template_path1: 第一个模板路径
            template_path2: 第二个模板路径
            
        Returns:
            差异对比结果
        """
        info1 = self.read_template(template_path1)
        info2 = self.read_template(template_path2)
        
        comparison = {
            'template1': info1['title'],
            'template2': info2['title'],
            'placeholder_diff': {
                'only_in_template1': list(set(info1['placeholders']) - set(info2['placeholders'])),
                'only_in_template2': list(set(info2['placeholders']) - set(info1['placeholders'])),
                'common': list(set(info1['placeholders']) & set(info2['placeholders']))
            },
            'structure_diff': {
                'paragraph_count': {
                    'template1': len(info1['paragraphs']),
                    'template2': len(info2['paragraphs'])
                },
                'table_count': {
                    'template1': len(info1['tables']),
                    'template2': len(info2['tables'])
                }
            }
        }
        
        return comparison


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(description='Word模板处理器')
    
    parser.add_argument('--action', choices=['read', 'fill', 'extract', 'batch', 'compare'],
                        required=True, help='操作类型')
    parser.add_argument('--template', type=str, help='模板文件路径')
    parser.add_argument('--template2', type=str, help='第二个模板路径（用于比较）')
    parser.add_argument('--data', type=str, help='填充数据（JSON格式字符串或文件路径）')
    parser.add_argument('--data_list', type=str, help='批量数据文件路径（JSON格式）')
    parser.add_argument('--output', type=str, help='输出文件路径')
    parser.add_argument('--output_dir', type=str, help='输出目录（批量生成时使用）')
    
    args = parser.parse_args()
    
    processor = WordTemplateProcessor()
    
    if args.action == 'read':
        if not args.template:
            print("错误：请提供 --template 参数")
            sys.exit(1)
        
        info = processor.read_template(args.template)
        print(json.dumps(info, ensure_ascii=False, indent=2))
    
    elif args.action == 'extract':
        if not args.template:
            print("错误：请提供 --template 参数")
            sys.exit(1)
        
        info = processor.extract_template_info(args.template, args.output)
        print(json.dumps(info, ensure_ascii=False, indent=2))
    
    elif args.action == 'fill':
        if not args.template or not args.data or not args.output:
            print("错误：请提供 --template、--data 和 --output 参数")
            sys.exit(1)
        
        # 解析数据
        if Path(args.data).exists():
            with open(args.data, 'r', encoding='utf-8') as f:
                data = json.load(f)
        else:
            data = json.loads(args.data)
        
        processor.fill_template(args.template, data, args.output)
    
    elif args.action == 'batch':
        if not args.template or not args.data_list or not args.output_dir:
            print("错误：请提供 --template、--data_list 和 --output_dir 参数")
            sys.exit(1)
        
        # 读取批量数据
        with open(args.data_list, 'r', encoding='utf-8') as f:
            data_list = json.load(f)
        
        processor.batch_fill(args.template, data_list, args.output_dir)
    
    elif args.action == 'compare':
        if not args.template or not args.template2:
            print("错误：请提供 --template 和 --template2 参数")
            sys.exit(1)
        
        comparison = processor.compare_templates(args.template, args.template2)
        print(json.dumps(comparison, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()

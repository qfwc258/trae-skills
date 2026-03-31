#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
法律文书生成器基础模块

提供统一的：
- 路径管理
- 占位符替换
- 错误处理
- 输出管理
- 校验功能
"""

import os
import json
import shutil
import logging
from datetime import datetime
from docx import Document


BLANK_PLACEHOLDER = '        '
CONFIG_FILE = 'config.json'
LOG_FILE = 'document_generator.log'


def setup_logging(log_file=None):
    """设置日志"""
    if log_file is None:
        log_file = os.path.join(get_script_dir(), LOG_FILE)

    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8'),
            logging.StreamHandler()
        ]
    )
    return logging.getLogger(__name__)


def load_config():
    """加载配置文件"""
    config_path = os.path.join(get_script_dir(), CONFIG_FILE)
    default_config = {
        'output_dir': 'output',
        'blank_placeholder': BLANK_PLACEHOLDER,
        'log_enabled': True
    }

    if os.path.exists(config_path):
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                return {**default_config, **config}
        except:
            pass
    return default_config


def save_config(config):
    """保存配置文件"""
    config_path = os.path.join(get_script_dir(), CONFIG_FILE)
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=2, ensure_ascii=False)


def get_script_dir():
    """获取脚本所在目录"""
    return os.path.dirname(os.path.abspath(__file__))


def get_project_dir():
    """获取项目根目录"""
    return os.path.dirname(get_script_dir())


def get_template_dir():
    """获取模板目录"""
    return os.path.join(get_project_dir(), 'legal_templates')


def get_output_dir():
    """获取输出目录（统一输出到 output/ 文件夹）"""
    config = load_config()
    output_path = os.path.join(get_project_dir(), config.get('output_dir', 'output'))

    if not os.path.exists(output_path):
        os.makedirs(output_path, exist_ok=True)

    return output_path


def get_log_dir():
    """获取日志目录"""
    return get_project_dir()


def copy_template_to_output(template_name, output_name):
    """
    复制模板到输出目录

    Args:
        template_name: 模板文件名
        output_name: 输出文件名

    Returns:
        output_path: 输出文件完整路径
    """
    template_path = os.path.join(get_template_dir(), template_name)
    output_dir = get_output_dir()
    output_path = os.path.join(output_dir, output_name)

    if not os.path.exists(template_path):
        raise FileNotFoundError(f'模板文件不存在: {template_path}')

    if os.path.exists(output_path):
        try:
            os.remove(output_path)
        except:
            name, ext = os.path.splitext(output_name)
            output_path = os.path.join(output_dir, f'{name}_{datetime.now().strftime("%Y%m%d%H%M%S")}{ext}')

    shutil.copy(template_path, output_path)
    return output_path


def replace_placeholders_in_doc(doc, placeholders_dict):
    """
    替换文档中的占位符

    Args:
        doc: Document 对象
        placeholders_dict: 占位符字典，格式 {占位符: 值}，值为空时替换为8个空格
    """
    for para in doc.paragraphs:
        full_text = para.text

        for placeholder, value in placeholders_dict.items():
            if placeholder in full_text:
                if value:
                    replace_value = value
                else:
                    replace_value = BLANK_PLACEHOLDER

                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replace_value)


def validate_placeholders_replaced(doc, placeholders_list):
    """
    校验文档中是否还有未替换的占位符

    Args:
        doc: Document 对象
        placeholders_list: 占位符列表

    Returns:
        (is_valid, remaining_placeholders): 是否有效，剩余占位符列表
    """
    remaining = []
    for para in doc.paragraphs:
        for placeholder in placeholders_list:
            if placeholder in para.text:
                remaining.append(placeholder)
    return len(remaining) == 0, remaining


def save_doc(doc, output_path):
    """保存文档"""
    doc.save(output_path)
    return output_path


def create_simple_document(template_name, output_name, placeholders_dict, validate=True):
    """
    创建简单文档的通用函数

    Args:
        template_name: 模板文件名
        output_name: 输出文件名
        placeholders_dict: 占位符字典
        validate: 是否校验占位符替换

    Returns:
        output_path: 输出文件路径
    """
    output_path = copy_template_to_output(template_name, output_name)
    doc = Document(output_path)
    replace_placeholders_in_doc(doc, placeholders_dict)
    save_doc(doc, output_path)

    if validate:
        config = load_config()
        placeholders = [k for k, v in placeholders_dict.items() if v == '']
        is_valid, remaining = validate_placeholders_replaced(doc, placeholders)
        if not is_valid:
            logging.warning(f'文档仍包含未替换的占位符: {remaining}')

    log_generation(template_name, output_name, placeholders_dict)
    return output_path


def log_generation(template_name, output_name, placeholders_dict):
    """记录生成日志"""
    config = load_config()
    if not config.get('log_enabled', True):
        return

    logger = setup_logging()
    log_entry = {
        'timestamp': datetime.now().isoformat(),
        'template': template_name,
        'output': output_name,
        'placeholders_filled': {k: v for k, v in placeholders_dict.items() if v}
    }
    logger.info(f'文档生成: {json.dumps(log_entry, ensure_ascii=False)}')


def get_generation_history(limit=50):
    """获取生成历史"""
    log_path = os.path.join(get_log_dir(), LOG_FILE)
    if not os.path.exists(log_path):
        return []

    history = []
    try:
        with open(log_path, 'r', encoding='utf-8') as f:
            for line in f:
                if '文档生成:' in line:
                    try:
                        json_str = line.split('文档生成:')[1].strip()
                        history.append(json.loads(json_str))
                    except:
                        continue
    except:
        pass

    return history[-limit:]


def create_batch_documents(template_name, output_base_name, data_list):
    """
    批量生成文档

    Args:
        template_name: 模板文件名
        output_base_name: 输出文件基础名
        data_list: 数据列表，每项为占位符字典

    Returns:
        output_paths: 生成的文档路径列表
    """
    output_dir = get_output_dir()
    output_paths = []

    for i, data in enumerate(data_list):
        output_name = f'{output_base_name}_{i+1}.docx'
        try:
            output_path = create_simple_document(template_name, output_name, data)
            output_paths.append(output_path)
            logging.info(f'批量生成 [{i+1}/{len(data_list)}]: {output_name}')
        except Exception as e:
            logging.error(f'批量生成失败 [{i+1}/{len(data_list)}]: {e}')

    return output_paths
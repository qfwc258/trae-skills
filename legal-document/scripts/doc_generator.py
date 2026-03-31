#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文书生成器
支持生成Word和PDF格式的法律文书
"""

import argparse
import sys
from pathlib import Path
from typing import Dict, List
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误：未安装 python-docx 库，请运行: pip install python-docx")
    sys.exit(1)

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.units import cm
except ImportError:
    print("警告：未安装 reportlab 库，PDF功能不可用。请运行: pip install reportlab")


class DocumentGenerator:
    """文书生成器"""
    
    def __init__(self):
        """初始化生成器"""
        self.word_supported = True
        self.pdf_supported = 'reportlab' in sys.modules
    
    def generate_word(self, content: str, output_path: str, metadata: Dict = None):
        """
        生成Word格式文书
        
        Args:
            content: 文书内容
            output_path: 输出文件路径
            metadata: 文书元数据（标题、当事人等）
        """
        if not self.word_supported:
            raise RuntimeError("Word生成功能不可用，请安装 python-docx")
        
        doc = Document()
        
        # 设置中文字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(12)
        
        # 添加标题
        if metadata and 'title' in metadata:
            title = doc.add_paragraph(metadata['title'])
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title_run.font.name = '黑体'
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            doc.add_paragraph()  # 空行
        
        # 处理内容
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # 判断是否为标题行（以【】或第X部分开头）
                is_heading = para_text.startswith('【') or para_text.startswith('第') or para_text.endswith('：')
                
                para = doc.add_paragraph(para_text)
                
                if is_heading:
                    para.runs[0].font.bold = True
                    para.runs[0].font.size = Pt(14)
        
        # 添加落款
        if metadata:
            doc.add_paragraph()  # 空行
            if 'date' in metadata:
                date_para = doc.add_paragraph(metadata.get('date', datetime.now().strftime('%Y年%m月%d日')))
                date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 保存文档
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        print(f"Word文档已生成: {output_path}")
        return str(output_path)
    
    def generate_pdf(self, content: str, output_path: str, metadata: Dict = None):
        """
        生成PDF格式文书
        
        Args:
            content: 文书内容
            output_path: 输出文件路径
            metadata: 文书元数据
        """
        if not self.pdf_supported:
            raise RuntimeError("PDF生成功能不可用，请安装 reportlab")
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 创建PDF文档
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        # 获取样式
        styles = getSampleStyleSheet()
        
        # 创建中文样式（需要注册中文字体）
        try:
            # 尝试注册常见中文字体
            font_paths = [
                '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
                '/usr/share/fonts/truetype/arphic/uming.ttc',
                'C:/Windows/Fonts/simsun.ttc',
                '/System/Library/Fonts/PingFang.ttc'
            ]
            
            font_registered = False
            for font_path in font_paths:
                if Path(font_path).exists():
                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    font_registered = True
                    break
            
            if font_registered:
                # 使用注册的中文字体创建样式
                title_style = ParagraphStyle(
                    'ChineseTitle',
                    parent=styles['Heading1'],
                    fontName='ChineseFont',
                    fontSize=18,
                    alignment=1,  # 居中
                    spaceAfter=20
                )
                
                normal_style = ParagraphStyle(
                    'ChineseNormal',
                    parent=styles['Normal'],
                    fontName='ChineseFont',
                    fontSize=12,
                    leading=20
                )
            else:
                raise Exception("未找到中文字体")
        
        except Exception as e:
            print(f"警告：{e}，使用默认字体（可能无法正确显示中文）")
            title_style = styles['Heading1']
            normal_style = styles['Normal']
        
        # 构建文档内容
        story = []
        
        # 添加标题
        if metadata and 'title' in metadata:
            story.append(Paragraph(metadata['title'], title_style))
            story.append(Spacer(1, 12))
        
        # 添加正文
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # 转义特殊字符
                para_text = para_text.replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(para_text, normal_style))
                story.append(Spacer(1, 6))
        
        # 添加日期
        if metadata and 'date' in metadata:
            story.append(Spacer(1, 20))
            story.append(Paragraph(metadata.get('date', ''), normal_style))
        
        # 生成PDF
        doc.build(story)
        
        print(f"PDF文档已生成: {output_path}")
        return str(output_path)
    
    def generate_from_template(self, template_data: Dict, output_path: str, format_type: str = 'word'):
        """
        基于模板数据生成文书
        
        Args:
            template_data: 模板数据（包含标题、内容、要素等）
            output_path: 输出路径
            format_type: 格式类型（word/pdf）
        """
        # 构建文书内容
        content = self._build_content(template_data)
        
        # 生成文档
        metadata = {
            'title': template_data.get('title', ''),
            'date': datetime.now().strftime('%Y年%m月%d日')
        }
        
        if format_type.lower() == 'word':
            return self.generate_word(content, output_path, metadata)
        elif format_type.lower() == 'pdf':
            return self.generate_pdf(content, output_path, metadata)
        else:
            raise ValueError(f"不支持的格式类型: {format_type}")
    
    def _build_content(self, template_data: Dict) -> str:
        """
        构建文书内容
        
        Args:
            template_data: 模板数据
            
        Returns:
            完整的文书内容
        """
        sections = []
        
        # 标题
        if 'title' in template_data:
            sections.append(f"【{template_data['title']}】\n")
        
        # 当事人信息
        if 'parties' in template_data:
            sections.append("【当事人信息】")
            for party in template_data['parties']:
                sections.append(f"{party['role']}: {party['name']}")
                if 'info' in party:
                    sections.append(f"  {party['info']}")
            sections.append("")
        
        # 主要内容
        for section in template_data.get('sections', []):
            section_title = section.get('title', '')
            section_content = section.get('content', '')
            
            if section_title:
                sections.append(f"【{section_title}】")
            
            if section_content:
                sections.append(section_content)
            
            sections.append("")
        
        # 法条引用
        if 'laws' in template_data:
            sections.append("【法律依据】")
            for law in template_data['laws']:
                sections.append(f"《{law['law_name']}》{law['article_number']}")
            sections.append("")
        
        # 落款
        if 'signature' in template_data:
            sections.append(template_data['signature'])
        
        return '\n'.join(sections)


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(description='法律文书生成器')
    
    parser.add_argument('--content', type=str, help='文书内容')
    parser.add_argument('--format', choices=['word', 'pdf'], default='word', help='输出格式')
    parser.add_argument('--output', type=str, required=True, help='输出文件路径')
    parser.add_argument('--title', type=str, help='文书标题')
    parser.add_argument('--template_data', type=str, help='模板数据文件路径（JSON格式）')
    
    args = parser.parse_args()
    
    generator = DocumentGenerator()
    
    # 从模板数据生成
    if args.template_data:
        import json
        with open(args.template_data, 'r', encoding='utf-8') as f:
            template_data = json.load(f)
        
        generator.generate_from_template(template_data, args.output, args.format)
    
    # 从内容生成
    elif args.content:
        metadata = {
            'title': args.title,
            'date': datetime.now().strftime('%Y年%m月%d日')
        }
        
        if args.format == 'word':
            generator.generate_word(args.content, args.output, metadata)
        else:
            generator.generate_pdf(args.content, args.output, metadata)
    
    else:
        print("请提供 --content 或 --template_data 参数")
        parser.print_help()
        sys.exit(1)


if __name__ == '__main__':
    main()

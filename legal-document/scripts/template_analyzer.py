#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模板分析器
深度分析Word模板的结构、格式、风格，即使没有占位符也能学习模板特征
"""

import argparse
import json
import sys
import re
from pathlib import Path
from typing import Dict, List, Any
from collections import Counter

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误：未安装 python-docx 库，请运行: pip install python-docx")
    sys.exit(1)


class TemplateAnalyzer:
    """模板分析器"""
    
    def __init__(self):
        """初始化分析器"""
        self.legal_terms = self._load_legal_terms()
    
    def _load_legal_terms(self) -> set:
        """加载法律术语库"""
        return {
            '原告', '被告', '第三人', '诉讼请求', '事实与理由', '证据',
            '起诉状', '答辩状', '上诉状', '申请书', '判决书', '裁定书',
            '当事人', '法定代表人', '委托代理人', '诉讼代理人',
            '合同', '协议', '条款', '违约', '赔偿', '责任',
            '甲方', '乙方', '丙方', '出租人', '承租人', '出借人', '借款人',
            '民法典', '合同法', '民事诉讼法', '刑法', '公司法',
            '判决', '裁定', '调解', '仲裁', '执行'
        }
    
    def analyze_template(self, template_path: str, output_json: str = None) -> Dict:
        """
        深度分析Word模板
        
        Args:
            template_path: 模板文件路径
            output_json: 输出JSON文件路径（可选）
            
        Returns:
            模板分析报告
        """
        template_path = Path(template_path)
        
        if not template_path.exists():
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        # 读取文档
        doc = Document(str(template_path))
        
        # 分析报告
        analysis = {
            'file_info': {
                'file_name': template_path.name,
                'file_path': str(template_path)
            },
            'document_structure': self._analyze_structure(doc),
            'format_styles': self._analyze_format_styles(doc),
            'language_style': self._analyze_language_style(doc),
            'content_pattern': self._analyze_content_pattern(doc),
            'sections': self._extract_sections(doc),
            'generation_guide': {}  # 给AI的生成指导
        }
        
        # 生成AI指导
        analysis['generation_guide'] = self._generate_ai_guide(analysis)
        
        # 保存分析报告
        if output_json:
            with open(output_json, 'w', encoding='utf-8') as f:
                json.dump(analysis, f, ensure_ascii=False, indent=2)
            print(f"模板分析报告已保存: {output_json}")
        
        return analysis
    
    def _analyze_structure(self, doc: Document) -> Dict:
        """
        分析文档结构
        
        Args:
            doc: Word文档对象
            
        Returns:
            结构分析结果
        """
        structure = {
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'title': '',
            'headings': [],
            'main_sections': []
        }
        
        # 提取标题（通常为第一个非空段落）
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                structure['title'] = text
                break
        
        # 识别标题和章节
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # 判断是否为标题行
            is_heading = self._is_heading(para, text)
            
            if is_heading:
                structure['headings'].append({
                    'index': i,
                    'text': text,
                    'level': self._determine_heading_level(text)
                })
                
                # 提取主要章节
                if self._is_main_section(text):
                    structure['main_sections'].append(text)
        
        return structure
    
    def _analyze_format_styles(self, doc: Document) -> Dict:
        """
        分析格式样式
        
        Args:
            doc: Word文档对象
            
        Returns:
            格式样式分析结果
        """
        styles = {
            'title_format': {},
            'heading_formats': [],
            'body_format': {},
            'fonts_used': set(),
            'font_sizes': []
        }
        
        # 分析标题格式
        if doc.paragraphs:
            first_para = doc.paragraphs[0]
            styles['title_format'] = self._extract_paragraph_format(first_para)
        
        # 分析各段落格式
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            format_info = self._extract_paragraph_format(para)
            
            # 收集字体信息
            if format_info.get('font_name'):
                styles['fonts_used'].add(format_info['font_name'])
            
            if format_info.get('font_size'):
                styles['font_sizes'].append(format_info['font_size'])
            
            # 记录标题格式
            if self._is_heading(para, para.text):
                styles['heading_formats'].append({
                    'text': para.text.strip()[:30],
                    'format': format_info
                })
        
        # 分析正文格式（取中间段落的平均值）
        body_paras = [p for p in doc.paragraphs[2:-2] if p.text.strip() and not self._is_heading(p, p.text)]
        if body_paras:
            styles['body_format'] = self._extract_paragraph_format(body_paras[0])
        
        # 转换集合为列表
        styles['fonts_used'] = list(styles['fonts_used'])
        styles['font_sizes'] = list(set(styles['font_sizes']))
        
        return styles
    
    def _analyze_language_style(self, doc: Document) -> Dict:
        """
        分析语言风格
        
        Args:
            doc: Word文档对象
            
        Returns:
            语言风格分析结果
        """
        all_text = '\n'.join([para.text for para in doc.paragraphs])
        
        style = {
            'legal_terms_used': [],
            'sentence_patterns': [],
            'formal_level': 'medium',
            'characteristics': []
        }
        
        # 提取使用的法律术语
        for term in self.legal_terms:
            if term in all_text:
                style['legal_terms_used'].append(term)
        
        # 分析句式模式
        sentences = re.split(r'[。！？\n]', all_text)
        sentence_starts = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 5:
                # 提取句首词
                start_words = sentence[:10]
                sentence_starts.append(start_words)
        
        # 识别典型句式
        patterns = self._identify_sentence_patterns(all_text)
        style['sentence_patterns'] = patterns
        
        # 判断正式程度
        if len(style['legal_terms_used']) > 10:
            style['formal_level'] = 'high'
            style['characteristics'].append('高度专业化的法律文书')
        elif len(style['legal_terms_used']) > 5:
            style['formal_level'] = 'medium'
            style['characteristics'].append('标准法律文书风格')
        else:
            style['formal_level'] = 'low'
            style['characteristics'].append('较为通俗的法律文书')
        
        return style
    
    def _analyze_content_pattern(self, doc: Document) -> Dict:
        """
        分析内容模式
        
        Args:
            doc: Word文档对象
            
        Returns:
            内容模式分析结果
        """
        pattern = {
            'has_parties_info': False,
            'has_requests': False,
            'has_facts': False,
            'has_evidence': False,
            'has_signature_area': False,
            'typical_elements': []
        }
        
        all_text = '\n'.join([para.text for para in doc.paragraphs])
        
        # 判断是否包含典型要素
        if any(kw in all_text for kw in ['原告', '被告', '甲方', '乙方', '出租人', '承租人']):
            pattern['has_parties_info'] = True
            pattern['typical_elements'].append('当事人信息')
        
        if any(kw in all_text for kw in ['诉讼请求', '请求事项', '合同条款']):
            pattern['has_requests'] = True
            pattern['typical_elements'].append('请求/条款')
        
        if any(kw in all_text for kw in ['事实', '理由', '经过', '背景']):
            pattern['has_facts'] = True
            pattern['typical_elements'].append('事实理由')
        
        if any(kw in all_text for kw in ['证据', '证明', '材料']):
            pattern['has_evidence'] = True
            pattern['typical_elements'].append('证据材料')
        
        if any(kw in all_text for kw in ['签名', '盖章', '日期', '年', '月', '日']):
            pattern['has_signature_area'] = True
            pattern['typical_elements'].append('落款签名')
        
        return pattern
    
    def _extract_sections(self, doc: Document) -> List[Dict]:
        """
        提取章节内容
        
        Args:
            doc: Word文档对象
            
        Returns:
            章节列表
        """
        sections = []
        current_section = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 判断是否为章节标题
            if self._is_main_section(text):
                # 保存上一个章节
                if current_section:
                    sections.append(current_section)
                
                # 开始新章节
                current_section = {
                    'title': text,
                    'content': [],
                    'paragraph_count': 0
                }
            else:
                # 添加到当前章节
                if current_section:
                    current_section['content'].append(text)
                    current_section['paragraph_count'] += 1
        
        # 保存最后一个章节
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _generate_ai_guide(self, analysis: Dict) -> Dict:
        """
        生成AI生成指导
        
        Args:
            analysis: 模板分析结果
            
        Returns:
            AI生成指导
        """
        guide = {
            'document_type': self._infer_document_type(analysis),
            'format_requirements': [],
            'structure_requirements': [],
            'language_requirements': [],
            'key_elements': []
        }
        
        # 格式要求
        title_format = analysis['format_styles'].get('title_format', {})
        if title_format.get('alignment') == 'center':
            guide['format_requirements'].append('标题居中对齐')
        
        if title_format.get('font_size'):
            guide['format_requirements'].append(f"标题字号：{title_format['font_size']}磅")
        
        # 结构要求
        for section in analysis['structure']['main_sections']:
            guide['structure_requirements'].append(f"包含章节：{section}")
        
        # 语言要求
        lang_style = analysis['language_style']
        if lang_style['legal_terms_used']:
            guide['language_requirements'].append(f"使用法律术语：{', '.join(lang_style['legal_terms_used'][:5])}")
        
        guide['language_requirements'].append(f"正式程度：{lang_style['formal_level']}")
        
        # 关键要素
        guide['key_elements'] = analysis['content_pattern']['typical_elements']
        
        return guide
    
    def _extract_paragraph_format(self, para) -> Dict:
        """提取段落格式信息"""
        format_info = {}
        
        try:
            # 对齐方式
            alignment = para.alignment
            if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                format_info['alignment'] = 'center'
            elif alignment == WD_ALIGN_PARAGRAPH.LEFT:
                format_info['alignment'] = 'left'
            elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                format_info['alignment'] = 'right'
            else:
                format_info['alignment'] = 'unknown'
            
            # 字体和字号
            if para.runs:
                run = para.runs[0]
                if run.font.name:
                    format_info['font_name'] = run.font.name
                if run.font.size:
                    format_info['font_size'] = run.font.size.pt
                
                # 是否加粗
                format_info['bold'] = run.font.bold or False
            
            # 行距（如果可获取）
            if para.paragraph_format.line_spacing:
                format_info['line_spacing'] = str(para.paragraph_format.line_spacing)
        
        except Exception as e:
            pass
        
        return format_info
    
    def _is_heading(self, para, text: str) -> bool:
        """判断是否为标题行"""
        # 基于格式判断
        if para.runs:
            run = para.runs[0]
            if run.font.bold or (run.font.size and run.font.size.pt > 14):
                return True
        
        # 基于内容判断
        heading_patterns = [
            r'^第[一二三四五六七八九十]+[条章节款]',
            r'^【.+】',
            r'^[一二三四五六七八九十]+[、.]',
            r'.+[:：]$',  # 以冒号结尾
            r'^(当事人|诉讼请求|事实与理由|证据|判决|裁定|合同|协议)'
        ]
        
        for pattern in heading_patterns:
            if re.match(pattern, text):
                return True
        
        # 段落较短且包含关键词
        if len(text) < 30 and any(kw in text for kw in ['原告', '被告', '诉讼请求', '事实', '理由', '证据', '合同', '条款']):
            return True
        
        return False
    
    def _determine_heading_level(self, text: str) -> int:
        """判断标题层级"""
        if re.match(r'^第[一二三四五六七八九十]+[章节]', text):
            return 1
        elif re.match(r'^第[一二三四五六七八九十]+条', text):
            return 2
        elif re.match(r'^[一二三四五六七八九十]+[、.]', text):
            return 3
        elif text.startswith('【'):
            return 2
        else:
            return 4
    
    def _is_main_section(self, text: str) -> bool:
        """判断是否为主要章节"""
        main_section_keywords = [
            '当事人', '诉讼请求', '请求事项', '事实与理由', '事实', '理由',
            '证据', '判决结果', '合同条款', '协议内容', '争议解决',
            '违约责任', '权利义务', '签署', '落款', '附则'
        ]
        
        return any(kw in text for kw in main_section_keywords)
    
    def _identify_sentence_patterns(self, text: str) -> List[str]:
        """识别句式模式"""
        patterns = []
        
        # 典型法律文书句式
        pattern_templates = [
            (r'根据.*规定', '引用法律条文'),
            (r'原告.*被告', '原被告表述'),
            (r'甲方.*乙方', '甲乙方表述'),
            (r'经.*协商', '协商表述'),
            (r'双方.*约定', '约定表述'),
            (r'依法.*判决', '判决表述'),
            (r'现.*如下', '说明表述')
        ]
        
        for pattern, name in pattern_templates:
            if re.search(pattern, text):
                patterns.append(name)
        
        return patterns
    
    def _infer_document_type(self, analysis: Dict) -> str:
        """推断文档类型"""
        title = analysis['document_structure'].get('title', '')
        
        if '起诉状' in title:
            return '诉讼文书-起诉状'
        elif '答辩状' in title:
            return '诉讼文书-答辩状'
        elif '合同' in title:
            return '合同协议'
        elif '申请书' in title:
            return '申请文书'
        elif '判决书' in title or '裁定书' in title:
            return '裁判文书'
        else:
            return '法律文书'


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(description='模板分析器 - 深度分析Word模板结构、格式、风格')
    
    parser.add_argument('--template', type=str, required=True, help='模板文件路径')
    parser.add_argument('--output', type=str, help='输出分析报告文件路径（JSON格式）')
    parser.add_argument('--brief', action='store_true', help='简化输出（仅显示摘要）')
    
    args = parser.parse_args()
    
    analyzer = TemplateAnalyzer()
    
    # 分析模板
    print(f"正在分析模板: {args.template}")
    print("=" * 80)
    
    analysis = analyzer.analyze_template(args.template, args.output)
    
    if args.brief:
        # 简化输出
        print(f"\n【文档类型】: {analysis['generation_guide']['document_type']}")
        print(f"\n【主要章节】:")
        for section in analysis['document_structure']['main_sections']:
            print(f"  - {section}")
        
        print(f"\n【关键要素】:")
        for element in analysis['content_pattern']['typical_elements']:
            print(f"  - {element}")
        
        print(f"\n【格式要求】:")
        for req in analysis['generation_guide']['format_requirements']:
            print(f"  - {req}")
        
        print(f"\n【语言要求】:")
        for req in analysis['generation_guide']['language_requirements']:
            print(f"  - {req}")
    
    else:
        # 完整输出
        print(json.dumps(analysis, ensure_ascii=False, indent=2))
    
    print(f"\n" + "=" * 80)
    print(f"分析完成！")
    
    if args.output:
        print(f"详细报告已保存至: {args.output}")


if __name__ == '__main__':
    main()

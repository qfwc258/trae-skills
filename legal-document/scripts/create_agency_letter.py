#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
使用 python-docx 生成代理词文档
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, Twips, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_color):
    """
    设置单元格背景色
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)


def add_formatted_paragraph(doc, text, font_name='仿宋_GB2312', font_size=14, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0, first_line=0, line_spacing=1.0):
    """
    添加格式化段落
    """
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if first_line > 0:
        para.paragraph_format.first_line_indent = Pt(first_line)
    
    # 设置行距
    para.paragraph_format.line_spacing = line_spacing
    para.paragraph_format.line_spacing_rule = 1  # 单倍行距

    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    return para


def create_agency_letter():
    """创建代理词文档"""
    doc = Document()

    # 设置页面
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(8.27)  # A4宽度
        section.page_height = Inches(11.69)  # A4高度
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1.18)
        section.top_margin = Inches(1.18)
        section.bottom_margin = Inches(1.18)

    # 添加标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run('代 理 词')
    run.font.name = '方正小标宋简体'
    run.font.size = Pt(22)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')

    # 添加称呼
    salutation = doc.add_paragraph()
    salutation.alignment = WD_ALIGN_PARAGRAPH.LEFT
    salutation.paragraph_format.space_before = Pt(12)
    salutation.paragraph_format.space_after = Pt(6)
    salutation.paragraph_format.first_line_indent = Pt(28)
    run = salutation.add_run('尊敬的审判长、审判员：')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 代理词正文
    opening = doc.add_paragraph()
    opening.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    opening.paragraph_format.space_before = Pt(6)
    opening.paragraph_format.space_after = Pt(6)
    opening.paragraph_format.first_line_indent = Pt(28)
    opening.paragraph_format.line_spacing = 1.0
    opening.paragraph_format.line_spacing_rule = 1
    run = opening.add_run('湖南金厚（宁乡）律师事务所依法接受本案原告湖南兴元科技股份有限公司（以下简称"原告"或"兴元公司"）的委托，指派我担任其与被告杭州仙果喔品牌运营管理有限公司（以下简称"被告"或"仙果喔公司"）承揽合同纠纷一案的诉讼代理人。通过庭审调查，结合全案证据，现就本案争议焦点发表如下代理意见，恳请合议庭予以采纳。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 一、主要观点
    section1_title = doc.add_paragraph()
    section1_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section1_title.paragraph_format.space_before = Pt(12)
    section1_title.paragraph_format.space_after = Pt(6)
    section1_title.paragraph_format.first_line_indent = Pt(0)
    run = section1_title.add_run('一、本案合同性质为典型的承揽合同，法律关系明确，应适用承揽合同的特殊规则。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    section1_content = doc.add_paragraph()
    section1_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section1_content.paragraph_format.space_before = Pt(6)
    section1_content.paragraph_format.space_after = Pt(6)
    section1_content.paragraph_format.first_line_indent = Pt(28)
    section1_content.paragraph_format.line_spacing = 1.0
    section1_content.paragraph_format.line_spacing_rule = 1
    run = section1_content.add_run('本案虽以《购销协议》为名，但究其实质，完全符合《中华人民共和国民法典》第七百七十条关于承揽合同的法律特征：')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 1. 工作成果的特定性与不可替代性
    point1_title = doc.add_paragraph()
    point1_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    point1_title.paragraph_format.space_before = Pt(6)
    point1_title.paragraph_format.space_after = Pt(3)
    point1_title.paragraph_format.first_line_indent = Pt(28)
    point1_title.paragraph_format.line_spacing = 1.0
    point1_title.paragraph_format.line_spacing_rule = 1
    run = point1_title.add_run('1. 工作成果的特定性与不可替代性：合同明确约定，标的物是"具有功能定制、外型定制的智能果昔机"（协议第1.3条），系为被告特定商业模式量身打造的非通用产品。这完全区别于买卖合同中可替代的通用商品。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 2. 定作人（被告）的主导地位
    point2_title = doc.add_paragraph()
    point2_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    point2_title.paragraph_format.space_before = Pt(3)
    point2_title.paragraph_format.space_after = Pt(3)
    point2_title.paragraph_format.first_line_indent = Pt(28)
    point2_title.paragraph_format.line_spacing = 1.0
    point2_title.paragraph_format.line_spacing_rule = 1
    run = point2_title.add_run('2. 定作人（被告）的主导地位：整个生产流程由被告深度控制与主导。被告不仅提供了自行研发的样机、技术构想（《产品配置清单》《产品性能需求表》），更通过派驻检验员全程驻厂监督（证据5《检验报告》），实际掌控了从技术标准到生产检验的全过程。原告仅作为"加工方"，严格按被告的指令和要求执行生产。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 3. 原告义务的本质是"完成工作"
    point3_title = doc.add_paragraph()
    point3_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    point3_title.paragraph_format.space_before = Pt(3)
    point3_title.paragraph_format.space_after = Pt(6)
    point3_title.paragraph_format.first_line_indent = Pt(28)
    point3_title.paragraph_format.line_spacing = 1.0
    point3_title.paragraph_format.line_spacing_rule = 1
    run = point3_title.add_run('3. 原告义务的本质是"完成工作"：原告的核心合同义务并非转移某一通用物的所有权，而是"按照双方签字盖章的样品……加工制作"（协议第1.2条），即利用自身生产能力，将被告的设计方案转化为有形的工作成果。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    section1_conclusion = doc.add_paragraph()
    section1_conclusion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section1_conclusion.paragraph_format.space_before = Pt(6)
    section1_conclusion.paragraph_format.space_after = Pt(12)
    section1_conclusion.paragraph_format.first_line_indent = Pt(28)
    section1_conclusion.paragraph_format.line_spacing = 1.0
    section1_conclusion.paragraph_format.line_spacing_rule = 1
    run = section1_conclusion.add_run('因此，本案定性为"承揽合同纠纷"准确无误。这意味着，本案的责任认定应严格遵循承揽合同"承揽人按定作人要求工作，定作人对设计及要求负责"的核心原则。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 二、第二部分
    section2_title = doc.add_paragraph()
    section2_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section2_title.paragraph_format.space_before = Pt(12)
    section2_title.paragraph_format.space_after = Pt(6)
    section2_title.paragraph_format.first_line_indent = Pt(0)
    run = section2_title.add_run('二、原告作为承揽人，已全面、精准地履行了合同核心义务——"按样生产"，工作成果合格。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    section2_content1 = doc.add_paragraph()
    section2_content1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section2_content1.paragraph_format.space_before = Pt(6)
    section2_content1.paragraph_format.space_after = Pt(3)
    section2_content1.paragraph_format.first_line_indent = Pt(28)
    section2_content1.paragraph_format.line_spacing = 1.0
    section2_content1.paragraph_format.line_spacing_rule = 1
    run = section2_content1.add_run('1. "按样生产"义务已履行完毕：原告的生产活动严格以双方共同签字封样的样机（证据2《样机封样照片》《样品承认书》）为唯一标准。所有量产产品在出厂前，均经过被告派驻人员的现场检验并书面确认合格（证据5《检验报告》）。这形成了完整的证据链，证明原告交付的工作成果完全符合被告作为定作人提出的"定制要求"。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    section2_content2 = doc.add_paragraph()
    section2_content2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section2_content2.paragraph_format.space_before = Pt(3)
    section2_content2.paragraph_format.space_after = Pt(12)
    section2_content2.paragraph_format.first_line_indent = Pt(28)
    section2_content2.paragraph_format.line_spacing = 1.0
    section2_content2.paragraph_format.line_spacing_rule = 1
    run = section2_content2.add_run('2. 工作成果已被定作人受领并投入商用，视为认可：被告不仅提取了已完成的299台设备，更关键的是，已将这些设备在全国范围内进行安装、调试并投入商业运营。根据《民法典》第七百八十条，定作人支付报酬的义务始于工作成果的交付与受领。被告将设备投入商业使用的行为，是其以事实行为对工作成果符合其使用目的的最终确认，其事后以"质量问题"否定先前行为，有违诚信。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 三、第三部分
    section3_title = doc.add_paragraph()
    section3_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section3_title.paragraph_format.space_before = Pt(12)
    section3_title.paragraph_format.space_after = Pt(6)
    section3_title.paragraph_format.first_line_indent = Pt(0)
    run = section3_title.add_run('三、被告作为定作人，拒付报酬的理由不能成立，其核心抗辩混淆了承揽合同中的责任划分。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    section3_opening = doc.add_paragraph()
    section3_opening.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section3_opening.paragraph_format.space_before = Pt(6)
    section3_opening.paragraph_format.space_after = Pt(3)
    section3_opening.paragraph_format.first_line_indent = Pt(28)
    section3_opening.paragraph_format.line_spacing = 1.0
    section3_opening.paragraph_format.line_spacing_rule = 1
    run = section3_opening.add_run('被告拒付报酬的主要理由是"产品质量问题"。然而，在承揽合同框架下，这一抗辩无法成立：')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    section3_point1 = doc.add_paragraph()
    section3_point1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section3_point1.paragraph_format.space_before = Pt(3)
    section3_point1.paragraph_format.space_after = Pt(3)
    section3_point1.paragraph_format.first_line_indent = Pt(28)
    section3_point1.paragraph_format.line_spacing = 1.0
    section3_point1.paragraph_format.line_spacing_rule = 1
    run = section3_point1.add_run('1. "质量问题"的根源在于定作人自身的设计缺陷，风险应由被告承担：这是本案最核心的争议点。被告自行提供的《质量问题列表》清晰显示，绝大多数故障（如结构密封不良导致进水、刀盘设计缺陷引震动、无排水设计致电机烧毁等）的根源，均被标注为"设计不合理"。根据《民法典》第七百七十六条，定作人提供图纸或技术要求不合理的，应当承担由此造成的损失。原告作为承揽人，其责任边界仅限于"是否严格按照被告提供的、可能存在缺陷的设计进行生产"。现有证据充分证明，原告已尽到此项义务。因此，因设计固有缺陷导致的产品性能风险，依法应由设计提供方即被告自行承担。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    section3_point2 = doc.add_paragraph()
    section3_point2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section3_point2.paragraph_format.space_before = Pt(3)
    section3_point2.paragraph_format.space_after = Pt(3)
    section3_point2.paragraph_format.first_line_indent = Pt(28)
    section3_point2.paragraph_format.line_spacing = 1.0
    section3_point2.paragraph_format.line_spacing_rule = 1
    run = section3_point2.add_run('2. 现有瑕疵属于"工作成果交付后的保修范畴"，不构成拒付报酬的合法抗辩：承揽合同中，定作人支付报酬的义务与工作成果的交付和受领直接挂钩。即使交付后的工作成果存在一些可修复的瑕疵，这属于合同约定的保修责任范围（参见《质量保证及售后服务协议》），应通过维修、更换部件等方式解决。这绝不构成定作人拒绝履行支付报酬这一主合同义务的法定理由。被告企图以局部、可修复的瑕疵来否定原告已完成全部定制工作的根本事实，并以此抵赖全部报酬，是典型的违约行为。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    section3_point3 = doc.add_paragraph()
    section3_point3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section3_point3.paragraph_format.space_before = Pt(3)
    section3_point3.paragraph_format.space_after = Pt(12)
    section3_point3.paragraph_format.first_line_indent = Pt(28)
    section3_point3.paragraph_format.line_spacing = 1.0
    section3_point3.paragraph_format.line_spacing_rule = 1
    run = section3_point3.add_run('3. 被告的行为已构成预期违约：被告不仅拒绝支付到期报酬，更明确表示将不再接受已为其生产完毕的剩余301台工作成果。该行为已明确表示其将不履行主要合同义务，构成《民法典》第五百七十八条规定的预期违约，原告有权提前追究其违约责任。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 四、第四部分
    section4_title = doc.add_paragraph()
    section4_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section4_title.paragraph_format.space_before = Pt(12)
    section4_title.paragraph_format.space_after = Pt(6)
    section4_title.paragraph_format.first_line_indent = Pt(0)
    run = section4_title.add_run('四、原告的诉讼请求具有坚实的事实基础与合同依据，且符合承揽合同的法律精神。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    section4_content1 = doc.add_paragraph()
    section4_content1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section4_content1.paragraph_format.space_before = Pt(6)
    section4_content1.paragraph_format.space_after = Pt(3)
    section4_content1.paragraph_format.first_line_indent = Pt(28)
    section4_content1.paragraph_format.line_spacing = 1.0
    section4_content1.paragraph_format.line_spacing_rule = 1
    run = section4_content1.add_run('1. 报酬请求权：原告主张的2,062,000元报酬，是基于被告未达到协议约定的2000台最低采购量而导致合同终止后，按约定的5800元/台结算标准计算得出（协议第2.2条）。这是被告未完成其承诺的采购量所应承担的商业后果，是原告完成工作后应得的对价。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    section4_content2 = doc.add_paragraph()
    section4_content2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section4_content2.paragraph_format.space_before = Pt(3)
    section4_content2.paragraph_format.space_after = Pt(3)
    section4_content2.paragraph_format.first_line_indent = Pt(28)
    section4_content2.paragraph_format.line_spacing = 1.0
    section4_content2.paragraph_format.line_spacing_rule = 1
    run = section4_content2.add_run('2. 逾期受领违约责任：协议第7.6条明确约定，被告逾期提货（即受领工作成果）需承担按日千分之一计算的违约金。该条款是针对定作人不及时受领工作成果的合法制约。被告长期拖延受领，导致原告资金被占用、库存积压，理应依约承担责任。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    section4_content3 = doc.add_paragraph()
    section4_content3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section4_content3.paragraph_format.space_before = Pt(3)
    section4_content3.paragraph_format.space_after = Pt(12)
    section4_content3.paragraph_format.first_line_indent = Pt(28)
    section4_content3.paragraph_format.line_spacing = 1.0
    section4_content3.paragraph_format.line_spacing_rule = 1
    run = section4_content3.add_run('3. 留置权与优先受偿权：根据《民法典》第七百八十三条，定作人未向承揽人支付报酬的，承揽人对完成的工作成果享有留置权。对于被告未支付报酬而原告合法占有的301台设备，原告依法享有留置权，并有权就该财产折价或拍卖、变卖后的价款优先受偿。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 五、第五部分
    section5_title = doc.add_paragraph()
    section5_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section5_title.paragraph_format.space_before = Pt(12)
    section5_title.paragraph_format.space_after = Pt(6)
    section5_title.paragraph_format.first_line_indent = Pt(0)
    run = section5_title.add_run('五、强调承揽合同定性对本案裁判的指导意义。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    section5_content = doc.add_paragraph()
    section5_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section5_content.paragraph_format.space_before = Pt(6)
    section5_content.paragraph_format.space_after = Pt(6)
    section5_content.paragraph_format.first_line_indent = Pt(28)
    section5_content.paragraph_format.line_spacing = 1.0
    section5_content.paragraph_format.line_spacing_rule = 1
    run = section5_content.add_run('将本案定性为承揽合同，其法律意义在于清晰界定了双方的风险与责任：')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 要点列表
    bullet1 = doc.add_paragraph()
    bullet1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bullet1.paragraph_format.space_before = Pt(3)
    bullet1.paragraph_format.space_after = Pt(3)
    bullet1.paragraph_format.first_line_indent = Pt(28)
    bullet1.paragraph_format.line_spacing = 1.0
    bullet1.paragraph_format.line_spacing_rule = 1
    run = bullet1.add_run('• 原告（承揽人）的责任是"按图施工"，风险在于施工过程是否符合指令。本案证据证明原告已完美履行。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    bullet2 = doc.add_paragraph()
    bullet2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bullet2.paragraph_format.space_before = Pt(3)
    bullet2.paragraph_format.space_after = Pt(12)
    bullet2.paragraph_format.first_line_indent = Pt(28)
    bullet2.paragraph_format.line_spacing = 1.0
    bullet2.paragraph_format.line_spacing_rule = 1
    run = bullet2.add_run('• 被告（定作人）的责任是"提供正确的图纸和要求"，并承担设计本身的风险。本案中产品的主要问题恰恰源于被告提供的"图纸"（设计）本身存在缺陷。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    section5_conclusion = doc.add_paragraph()
    section5_conclusion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    section5_conclusion.paragraph_format.space_before = Pt(6)
    section5_conclusion.paragraph_format.space_after = Pt(12)
    section5_conclusion.paragraph_format.first_line_indent = Pt(28)
    section5_conclusion.paragraph_format.line_spacing = 1.0
    section5_conclusion.paragraph_format.line_spacing_rule = 1
    run = section5_conclusion.add_run('因此，被告将其自身设计缺陷所引发的商业风险，试图转化为对承揽人的索赔主张，完全颠倒了承揽合同中的基本责任逻辑，不应获得法律支持。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 结论
    conclusion = doc.add_paragraph()
    conclusion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    conclusion.paragraph_format.space_before = Pt(12)
    conclusion.paragraph_format.space_after = Pt(6)
    conclusion.paragraph_format.first_line_indent = Pt(28)
    conclusion.paragraph_format.line_spacing = 1.0
    conclusion.paragraph_format.line_spacing_rule = 1
    run = conclusion.add_run('综上所述，原告已忠实、完整地履行了承揽人的合同义务，交付了符合定作人特定要求的工作成果。被告作为定作人，在受领并使用了工作成果后，不仅不履行支付报酬的核心义务，反而将自身设计责任导致的后果归咎于原告，其行为已构成根本违约。原告的诉讼请求事实清楚、证据确凿、于法有据。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    appeal = doc.add_paragraph()
    appeal.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    appeal.paragraph_format.space_before = Pt(6)
    appeal.paragraph_format.space_after = Pt(12)
    appeal.paragraph_format.first_line_indent = Pt(28)
    appeal.paragraph_format.line_spacing = 1.0
    appeal.paragraph_format.line_spacing_rule = 1
    run = appeal.add_run('恳请贵院明察，依法支持原告的全部诉讼请求，以维护公平诚信的市场交易秩序。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 此致
    ending1 = doc.add_paragraph()
    ending1.alignment =
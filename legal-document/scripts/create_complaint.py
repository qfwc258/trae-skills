#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成民事起诉状
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_complaint():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(script_dir)
    output_dir = os.getcwd()

    doc = Document()

    sections = doc.sections
    for section in sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1.18)
        section.top_margin = Inches(1.18)
        section.bottom_margin = Inches(1.18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run('民 事 起 诉 状')
    run.font.name = '方正小标宋简体'
    run.font.size = Pt(22)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')

    section_title = doc.add_paragraph()
    section_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    section_title.paragraph_format.space_before = Pt(12)
    section_title.paragraph_format.space_after = Pt(6)
    run = section_title.add_run('原告：')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    plaintiff_info = doc.add_paragraph()
    plaintiff_info.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    plaintiff_info.paragraph_format.space_before = Pt(0)
    plaintiff_info.paragraph_format.space_after = Pt(6)
    plaintiff_info.paragraph_format.first_line_indent = Pt(28)
    run = plaintiff_info.add_run('李东明，男，1972年6月19日出生，汉族，住山东省泰安市岱岳区马庄镇南王村100号。公民身份号码：37091119720619567X。联系电话：18108456677。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    defendant1_title = doc.add_paragraph()
    defendant1_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    defendant1_title.paragraph_format.space_before = Pt(12)
    defendant1_title.paragraph_format.space_after = Pt(6)
    run = defendant1_title.add_run('被告一：')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    defendant1_info = doc.add_paragraph()
    defendant1_info.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    defendant1_info.paragraph_format.space_before = Pt(0)
    defendant1_info.paragraph_format.space_after = Pt(6)
    defendant1_info.paragraph_format.first_line_indent = Pt(28)
    run = defendant1_info.add_run('湖南玟鹰建筑劳务有限公司，住所地湖南省长沙市宁乡高新技术产业园区新天北路001号综合楼众创空间40680号，统一社会信用代码：91430124MACPX2H56Q。联系电话：15874571333。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    defendant1_rep = doc.add_paragraph()
    defendant1_rep.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    defendant1_rep.paragraph_format.space_before = Pt(0)
    defendant1_rep.paragraph_format.space_after = Pt(6)
    defendant1_rep.paragraph_format.first_line_indent = Pt(28)
    run = defendant1_rep.add_run('法定代表人：胡小宏。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    defendant2_title = doc.add_paragraph()
    defendant2_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    defendant2_title.paragraph_format.space_before = Pt(12)
    defendant2_title.paragraph_format.space_after = Pt(6)
    run = defendant2_title.add_run('被告二：')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    defendant2_info = doc.add_paragraph()
    defendant2_info.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    defendant2_info.paragraph_format.space_before = Pt(0)
    defendant2_info.paragraph_format.space_after = Pt(6)
    defendant2_info.paragraph_format.first_line_indent = Pt(28)
    run = defendant2_info.add_run('湖南浩澜建筑工程有限公司，住所地：湖南省长沙市宁乡高新技术产业园区新天北路001号湖南省大学科技产业园综合楼4楼40661号，统一社会信用代码：91430181MA4TD6NX4G。联系电话：15074903572。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    defendant2_rep = doc.add_paragraph()
    defendant2_rep.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    defendant2_rep.paragraph_format.space_before = Pt(0)
    defendant2_rep.paragraph_format.space_after = Pt(6)
    defendant2_rep.paragraph_format.first_line_indent = Pt(28)
    run = defendant2_rep.add_run('法定代表人：梁云。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    request_title = doc.add_paragraph()
    request_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    request_title.paragraph_format.space_before = Pt(12)
    request_title.paragraph_format.space_after = Pt(6)
    run = request_title.add_run('诉讼请求：')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    request1 = doc.add_paragraph()
    request1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    request1.paragraph_format.space_before = Pt(0)
    request1.paragraph_format.space_after = Pt(6)
    request1.paragraph_format.first_line_indent = Pt(28)
    run = request1.add_run('一、请求判决被告一、被告二向原告支付拖欠的劳务报酬36435元及资金占用利息2407元（利息自2024年1月17日起暂计算至2025年12月15日，后续按年利率3.45%支付至实际清偿之日止）；')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    request2 = doc.add_paragraph()
    request2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    request2.paragraph_format.space_before = Pt(0)
    request2.paragraph_format.space_after = Pt(12)
    request2.paragraph_format.first_line_indent = Pt(28)
    run = request2.add_run('二、本案诉讼费用由两被告负担。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    fact_title = doc.add_paragraph()
    fact_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fact_title.paragraph_format.space_before = Pt(12)
    fact_title.paragraph_format.space_after = Pt(6)
    run = fact_title.add_run('事实与理由：')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    fact1 = doc.add_paragraph()
    fact1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fact1.paragraph_format.space_before = Pt(0)
    fact1.paragraph_format.space_after = Pt(6)
    fact1.paragraph_format.first_line_indent = Pt(28)
    run = fact1.add_run('2023年8月至2024年1月16日，原告在位于湖南省宁乡市的"长沙IVD生产基地建设项目"工地从事植筋工作。原告工作经班组长李东春（已故）安排，纳入其劳务班组。该工程总承包人为被告二湖南浩澜建筑工程有限公司，劳务分包单位为被告一湖南玟鹰建筑劳务有限公司。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    fact2 = doc.add_paragraph()
    fact2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fact2.paragraph_format.space_before = Pt(0)
    fact2.paragraph_format.space_after = Pt(6)
    fact2.paragraph_format.first_line_indent = Pt(28)
    run = fact2.add_run('项目完工后，经核算，确认被告拖欠原告的劳务报酬共计36435元。原告多次向被告一负责人胡小宏追索，但均未果。被告一甚至在班组长李东春去世后，否认原告曾在其工地工作的事实。原告亦曾寻求劳动行政部门调解，但未能解决纠纷。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    fact3 = doc.add_paragraph()
    fact3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fact3.paragraph_format.space_before = Pt(0)
    fact3.paragraph_format.space_after = Pt(12)
    fact3.paragraph_format.first_line_indent = Pt(28)
    run = fact3.add_run('原告认为，依据《保障农民工工资支付条例》等规定，被告一作为直接用工单位、被告二作为总承包单位，对拖欠原告的劳务报酬应承担清偿责任。为维护原告的合法权益，特向贵院提起诉讼，望判如所请。')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    ending1 = doc.add_paragraph()
    ending1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ending1.paragraph_format.space_before = Pt(12)
    ending1.paragraph_format.space_after = Pt(6)
    run = ending1.add_run('此致')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    court = doc.add_paragraph()
    court.alignment = WD_ALIGN_PARAGRAPH.LEFT
    court.paragraph_format.space_before = Pt(6)
    court.paragraph_format.space_after = Pt(24)
    run = court.add_run('宁乡市人民法院')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    sign_title = doc.add_paragraph()
    sign_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sign_title.paragraph_format.space_before = Pt(6)
    sign_title.paragraph_format.space_after = Pt(6)
    run = sign_title.add_run('具状人：')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_para.paragraph_format.space_before = Pt(6)
    date_para.paragraph_format.space_after = Pt(6)
    run = date_para.add_run('年    月    日')
    run.font.name = '仿宋_GB2312'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    output_path = os.path.join(output_dir, '民事起诉状_李东明vs玟鹰浩澜_20260322.docx')
    doc.save(output_path)
    print(f'民事起诉状已生成: {output_path}')


if __name__ == '__main__':
    create_complaint()

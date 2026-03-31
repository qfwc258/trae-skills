#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
证据目录生成器 - 支持填充表格数据

用法:
    python create_evidence_directory_full.py
"""

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import shutil
import os


def set_cell_font(cell, font_name='仿宋_GB2312', font_size=12, center_align=False):
    """设置单元格字体格式"""
    for paragraph in cell.paragraphs:
        if center_align:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = font_name
            r = run._element
            rFonts = r.get_or_add_rPr().get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)


def create_evidence_directory_full():
    """生成填充完整数据的证据目录"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(script_dir)
    template_path = os.path.join(project_dir, 'legal_templates', '3 证据目录 .docx')
    output_dir = os.getcwd()
    output_path = os.path.join(output_dir, '证据目录_易良波_建设工程施工合同纠纷.docx')

    if os.path.exists(output_path):
        try:
            os.remove(output_path)
        except:
            output_path = os.path.join(output_dir, '证据目录_易良波_建设工程施工合同纠纷_new.docx')

    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    for para in doc.paragraphs:
        full_text = para.text
        if 'weitr' in full_text:
            for run in para.runs:
                if 'weitr' in run.text:
                    run.text = run.text.replace('weitr', '易良波')

    table = doc.tables[0]

    evidence_data = [
        ("1", "《装修油漆改造工程施工合同》", "1-2", '1. 证明原告与被告一（贵州凌宇建筑工程有限公司潇湘分公司）之间存在合法有效的建设工程施工合同关系。\n2. 证明合同约定了工程内容、工期、价款、付款方式等。\n3. 证明合同第十条明确约定了履约保证金的交纳与返还条件（"乙方进场三天交满履约保证金。乙方开工一个月付款时一次性退还"）。', "原告"),
        ("2", "微信支付转账电子凭证\n（转账单号：1000050001202506010628961068476）", "3", '证明原告已于2025年6月1日，按合同约定向被告一指定的收款人"邹峰"支付了履约保证金人民币30,000元。', "原告"),
        ("3", "微信支付转账电子凭证\n（转账单号：1000050001202506011328005469954）", "4", "证明原告为促成涉案合同的签订，于2025年6月1日向中介方支付了中介费人民币3,000元。", "原告"),
        ("4", "《不可撤销劳务信息咨询服务支付合同》", "4-5", '1. 证明原告为获取涉案工程项目，与第三方中介签订了居间合同。\n2. 证明该笔3,000元中介费的支付与涉案合同直接相关，是因被告一违约而造成的实际损失。', "原告"),
        ("5", "《易良波与洪亮电话录音（文字版）》", "5-6", '1. 证明被告一的工作人员洪亮承认项目因被告自身原因（"开工令"下达延迟）导致延误。\n2. 证明被告一承认欠款事实，并明确承诺于2025年11月24日（星期一）中午前归还"3万3"元（即保证金及部分损失）。\n3. 证明被告一后续未履行其承诺，构成违约。', "原告"),
        ("6", "国家企业信用信息公示系统报告\n（贵州凌宇建筑工程有限公司潇湘分公司）", "6-8", '1. 证明被告一的主体资格、负责人、住所地等信息。\n2. 证明被告一为"有限责任公司分公司"，不具有独立法人资格。\n3. 证明被告一已被列入经营异常名录，其经营状况异常。', "公开信息"),
        ("7", "国家企业信用信息公示系统报告\n（贵州凌宇建筑工程有限公司）", "9-11", '1. 证明被告二的主体资格、法定代表人、住所地等信息。\n2. 证明被告二为依法设立的有限责任公司，是被告一的法人（总公司）。\n3. 为要求被告二对被告一的债务承担连带责任提供依据。', "公开信息"),
        ("8", "原告《营业执照》复印件", "12", "证明原告长沙易湘建筑劳务有限公司的主体资格及诉讼主体资格。", "原告"),
        ("9", "《法定代表人身份证明书》", "13", "证明易良波系原告公司的法定代表人，其有权代表公司进行诉讼活动。", "原告"),
    ]

    for i, (num, name, page, purpose, source) in enumerate(evidence_data):
        if i + 1 < len(table.rows):
            row = table.rows[i + 1]
            row.cells[0].text = num
            row.cells[1].text = name
            row.cells[2].text = page
            row.cells[3].text = purpose
            row.cells[4].text = source

            set_cell_font(row.cells[0], center_align=True)
            set_cell_font(row.cells[1])
            set_cell_font(row.cells[2], center_align=True)
            set_cell_font(row.cells[3])
            set_cell_font(row.cells[4])

    doc.save(output_path)
    print(f'✅ 证据目录已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    create_evidence_directory_full()
